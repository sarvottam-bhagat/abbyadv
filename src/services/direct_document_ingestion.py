"""Fast ingestion path for chat attachments with extractable text."""
from io import BytesIO
from pathlib import Path

from sqlalchemy.ext.asyncio import AsyncSession

from src.database.models import CaseDocument
from src.services.document_pipeline import index_document
from src.services.storage import StorageService


def extract_text(content: bytes, file_name: str) -> str:
    suffix = Path(file_name).suffix.lower()
    if suffix in {".txt", ".md", ".csv", ".json"}:
        return content.decode("utf-8", errors="ignore")
    if suffix == ".pdf":
        from pypdf import PdfReader
        reader = PdfReader(BytesIO(content))
        return "\n\n".join(page.extract_text() or "" for page in reader.pages)
    if suffix == ".docx":
        from docx import Document
        document = Document(BytesIO(content))
        paragraphs = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
        for table in document.tables:
            paragraphs.extend(" | ".join(cell.text.strip() for cell in row.cells) for row in table.rows)
        return "\n".join(paragraphs)
    raise ValueError("Unsupported attachment. Supported formats: PDF, DOCX, TXT, MD, CSV, JSON")


async def ingest_chat_attachment(db: AsyncSession, document: CaseDocument) -> int:
    content = await StorageService().download(document.storage_key)
    text = extract_text(content, document.file_name).strip()
    if not text:
        raise ValueError("No text could be extracted. Scanned PDFs require ABBYY processing.")
    document.extracted_text = text[:200_000]
    document.processing_status = "processed"
    document.ocr_status = "processed"
    document.error_message = None
    await db.commit()
    return await index_document(db, document)


async def read_chat_attachment(db: AsyncSession, document: CaseDocument) -> str:
    """Extract a one-off chat attachment without creating Qdrant vectors."""
    content = await StorageService().download(document.storage_key)
    text = extract_text(content, document.file_name).strip()
    if not text:
        raise ValueError("No text could be extracted. Scanned PDFs require ABBYY processing.")
    document.extracted_text = text[:200_000]
    document.processing_status = "processed"
    document.ocr_status = "processed"
    document.embedding_status = "not_requested"
    document.error_message = None
    await db.commit()
    return document.extracted_text
