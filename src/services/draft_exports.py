"""Professional, on-demand exports for saved AbbyAdv drafts."""
from html import escape
from io import BytesIO
import re


def _lines(content: str) -> list[str]:
    return [line.strip() for line in content.replace("\r\n", "\n").split("\n")]


def _plain(value: str) -> str:
    return re.sub(r"([*_`#])", "", value).strip()


def build_docx(title: str, content: str) -> bytes:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt

    document = Document(); section = document.sections[0]
    section.top_margin = section.bottom_margin = Inches(0.8)
    section.left_margin = section.right_margin = Inches(0.9)
    normal = document.styles["Normal"]; normal.font.name = "Times New Roman"; normal.font.size = Pt(11)
    heading = document.add_paragraph(); heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = heading.add_run(title.upper()); run.bold = True; run.font.name = "Times New Roman"; run.font.size = Pt(13)
    for line in _lines(content):
        if not line:
            document.add_paragraph(); continue
        text = _plain(line)
        if line.startswith("### ") or line.startswith("## ") or line.startswith("# "):
            paragraph = document.add_paragraph(); run = paragraph.add_run(text); run.bold = True; run.font.name = "Times New Roman"; run.font.size = Pt(11)
        else:
            paragraph = document.add_paragraph(text); paragraph.paragraph_format.space_after = Pt(7); paragraph.paragraph_format.line_spacing = 1.15
    output = BytesIO(); document.save(output); return output.getvalue()


def build_pdf(title: str, content: str) -> bytes:
    from reportlab.lib.enums import TA_CENTER
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import inch
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

    output = BytesIO(); document = SimpleDocTemplate(output, pagesize=A4, leftMargin=0.85*inch, rightMargin=0.85*inch, topMargin=0.75*inch, bottomMargin=0.75*inch)
    styles = getSampleStyleSheet(); title_style = ParagraphStyle("LegalTitle", parent=styles["Title"], fontName="Times-Bold", fontSize=13, leading=17, alignment=TA_CENTER, spaceAfter=16)
    heading_style = ParagraphStyle("LegalHeading", parent=styles["Heading2"], fontName="Times-Bold", fontSize=11, leading=14, spaceBefore=10, spaceAfter=6)
    body_style = ParagraphStyle("LegalBody", parent=styles["BodyText"], fontName="Times-Roman", fontSize=10.5, leading=15, spaceAfter=8)
    story = [Paragraph(escape(title.upper()), title_style)]
    for line in _lines(content):
        if not line:
            story.append(Spacer(1, 5)); continue
        text = escape(_plain(line)).replace("\n", "<br/>")
        story.append(Paragraph(text, heading_style if line.startswith("#") else body_style))
    document.build(story); return output.getvalue()
